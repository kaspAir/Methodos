"""Beweist: HERMES PIA weist je Abschnitt Herkunft + Begründung aus (Nachweis-Anhang)."""
from io import BytesIO

from docx import Document

from app.config import get_config
from app.domains.catalog.service import CatalogService
from app.domains.generation.service import GenerationService
from app.domains.interview.service import InterviewService
from app.domains.method.service import MethodService


def _interview():
    cfg = get_config()
    return InterviewService(MethodService(cfg.METHODS_DIR), CatalogService(cfg.CATALOGS_DIR), None)


class _Sess:
    method_id = "hermes_pia"
    project_name = "Testprojekt"
    project_type_id = "fachanwendung_einfuehrung"
    auftraggeber = "Amt X"


def test_herkunft_deterministisch():
    assert InterviewService._herkunft("PL hat gesprochen", []) == "Projektleiter (Interview)"
    assert InterviewService._herkunft("", []) == "HERMES PIA (kombiniert)"
    mixed = InterviewService._herkunft("PL Text", [{"type": "offer", "status": "accepted"}])
    assert mixed == "Projektleiter + HERMES PIA"


def test_inhalt_summary_tabelle_und_text():
    assert InterviewService._inhalt_summary({"text": "Hallo Welt"}) == "Hallo Welt"
    s = InterviewService._inhalt_summary([{"ergebnis": "Studie"}, {"ergebnis": "Prototyp"}])
    assert "Studie" in s and "Prototyp" in s


def test_build_nachweis_ohne_llm_nutzt_fallback():
    svc = _interview()
    answers = {
        "ausgangslage": {"raw_text": "Das SAP-System ist zu teuer.",
                          "extracted": {"text": "Das bestehende SAP-System ist kostenintensiv."},
                          "followups": []},
        "termine": {"raw_text": "",
                    "extracted": [{"ergebnis": "Studie", "abnahme": "Projektleiter"}],
                    "followups": []},
    }
    nw = svc.build_nachweis(_Sess(), answers)
    by_abschnitt = {n["abschnitt"]: n for n in nw}
    assert by_abschnitt["Ausgangslage"]["herkunft"] == "Projektleiter (Interview)"
    assert by_abschnitt["Ergebnisse und Termine"]["herkunft"] == "HERMES PIA (kombiniert)"
    # Begründung ist gesetzt (deterministischer Fallback ohne LLM)
    assert all(n["begruendung"] for n in nw)


def test_leere_abschnitte_kommen_nicht_in_den_nachweis():
    svc = _interview()
    answers = {"sachmittel": {"raw_text": "", "extracted": [], "followups": []}}
    assert svc.build_nachweis(_Sess(), answers) == []


def test_generierung_haengt_nachweis_anhang_an():
    cfg = get_config()
    gen = GenerationService(MethodService(cfg.METHODS_DIR))
    metadata = {"projektname": "Testprojekt", "version": "0.1"}
    answers = {"ausgangslage": {"raw_text": "x", "extracted": {"text": "Eine Ausgangslage."}}}
    nachweis = [
        {"abschnitt": "Ausgangslage", "herkunft": "Projektleiter (Interview)",
         "begruendung": "Beruht auf den Angaben des Projektleiters."},
    ]
    buf = gen.generate("hermes_pia", answers, metadata, nachweis=nachweis)
    assert isinstance(buf, BytesIO)
    doc = Document(buf)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Nachweis" in full
    # Anhang-Tabelle enthält Herkunft + Begründung
    found = False
    for t in doc.tables:
        for row in t.rows:
            joined = " ".join(c.text for c in row.cells)
            if "Projektleiter (Interview)" in joined and "Angaben des Projektleiters" in joined:
                found = True
    assert found, "Nachweis-Tabelle nicht gefunden"
