"""Beweist: Enddatum gefüllt, 'Projektauftrag' korrigiert, Kap.-6-Spalten passen."""
from docx import Document

from app.config import get_config
from app.domains.generation.service import GenerationService, W, _max_termin
from app.domains.method.service import MethodService


def _gen():
    cfg = get_config()
    return GenerationService(MethodService(cfg.METHODS_DIR))


def _full_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_max_termin():
    rows = [{"termin": "05.10.2026"}, {"termin": "16.11.2026"}, {"termin": "01.09.2026"}]
    assert _max_termin(rows) == "16.11.2026"


def test_enddatum_wird_gefuellt():
    gen = _gen()
    answers = {"termine": {"extracted": [
        {"ergebnis": "Stakeholder-Liste", "termin": "28.09.2026"},
        {"ergebnis": "Meilenstein Durchfuehrungsfreigabe", "termin": "16.11.2026"},
    ]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    full = _full_text(doc)
    assert "16.11.2026" in full
    assert "tt.mm.jjjj" not in full.split("Enddatum")[-1][:60]


def test_projektauftrag_wird_zu_durchfuehrungsauftrag():
    gen = _gen()
    answers = {"risiken": {"extracted": [
        {"beschreibung": "Risiko", "ew": "Mittel", "ag": "Hoch",
         "massnahmen": "Im Projektauftrag verankern", "verantwortung": "Projektleiter", "termin": "laufend"},
    ]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    full_cells = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "Durchführungsauftrag verankern" in full_cells
    assert "Projektauftrag" not in full_cells


def test_projektorganisation_bestaetigung_in_letzter_spalte():
    gen = _gen()
    row = {"rolle_person": "Projektleiter", "bestaetigung": "ausstehend"}
    for i in range(1, 10):
        row[f"monat_{i}"] = "5" if i <= 3 else ""
    answers = {"projektorganisation": {"extracted": [row]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    W_TR = f"{{{W}}}tr"
    for tbl in doc.element.body.iter(f"{{{W}}}tbl"):
        if "Bestätigung" in "".join(tbl.itertext()) and "Monat 1" in "".join(tbl.itertext()):
            # Datenzeile finden: letzte Zelle muss 'ausstehend' sein, Monat 1 = '5'
            from docx.table import Table
            t = Table(tbl, doc)
            data = [r for r in t.rows if "Projektleiter" in r.cells[0].text]
            assert data, "Datenzeile nicht gefunden"
            cells = [c.text.strip() for c in data[0].cells]
            assert cells[-1] == "ausstehend", cells
            assert cells[1] == "5"  # Monat 1
            break
    else:
        raise AssertionError("Projektorganisation-Tabelle nicht gefunden")
